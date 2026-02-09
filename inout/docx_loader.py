from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document

@dataclass(frozen=True, slots=True)
class DocxLoader():
    """
    Loads paragraph text from a .docx file
    """

    strip_whitespace: bool = True
    keep_empty_paragraphs: bool = True

    def load_paragraphs(self, docx_path: str | Path) -> list[str]:
        """
        Read a .docx and return a list of paragraph strings in document order.
        
        Parameters
        ----------
        docx_path:
            Path (or string path) to a .docx file.

        Returns
        ----------
        list[str]
            Paragraph texts, optionally stripped and optionally excluding empties
        """
        path = Path(docx_path)
        self._validate_docx_path(path)

        doc = Document(str(path))
        paras = [p.text for p in doc.paragraphs]
        return self._postprocess(paras)
    
    def iter_paragraphs(self, docx_path: str | Path) -> Iterable[str]:
        """
        Generator version of `load_paragraphs`. Useful for streaming.
        """
        path = Path(docx_path)
        self._validate_docx_path(path)

        doc = Document(str(path))
        for p in doc.paragraphs:
            txt = p.text
            if self.strip_whitespace:
                txt = txt.strip()
            if not self.keep_empty_paragraphs and not txt:
                continue
            yield txt

    def _postprocess(self, paragraphs: list[str]) -> list[str]:
        if self.strip_whitespace:
            paragraphs = [p.strip() for p in paragraphs]
        if not self.keep_empty_paragraphs:
            paragraphs = [p for p in paragraphs if p]
        return paragraphs

    @staticmethod
    def _validate_docx_path(path: Path) -> None:
        if not path.exists():
            raise FileNotFoundError(f"Docx file not found: {path}")
        if not path.is_file():
            raise ValueError(f"Docx path is not a file: {path}")
        if path.suffix.lower() != ".docx":
            raise ValueError(f"Expected a .docx file, but got: {path.name}")
