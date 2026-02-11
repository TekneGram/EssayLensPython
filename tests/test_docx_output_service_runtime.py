from __future__ import annotations

import unittest
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document
from services.docx_output_service import DocxOutputService


class DocxOutputServiceRuntimeTests(unittest.TestCase):
    def test_post_init_constructs_editor_with_author(self) -> None:
        with patch("services.docx_output_service.TrackChangesEditor") as editor_cls:
            svc = DocxOutputService(author="Alice")
        editor_cls.assert_called_once_with(author="Alice")
        self.assertIsNotNone(svc)

    def test_build_report_with_header_and_body_forwards_arguments(self) -> None:
        with patch("services.docx_output_service.TrackChangesEditor") as editor_cls:
            editor = MagicMock()
            editor_cls.return_value = editor
            svc = DocxOutputService(author="Alice")

            svc.build_report_with_header_and_body(
                output_path=Path("out.docx"),
                original_paragraphs=["p1", "p2"],
                edited_text="edited full",
                header_lines=["Name: Dan", "Course: X"],
                edited_body_text="edited body",
                corrected_body_text="corrected body",
                feedback_paragraphs=["Good", "## Next"],
                include_edited_text=False,
            )

            editor.build_report_with_header_and_body.assert_called_once_with(
                output_path="out.docx",
                original_paragraphs=["p1", "p2"],
                edited_text="edited full",
                header_lines=["Name: Dan", "Course: X"],
                edited_body_text="edited body",
                corrected_body_text="corrected body",
                feedback_heading="Language Feedback",
                feedback_paragraphs=["Good", "## Next"],
                feedback_as_tracked_insertion=False,
                add_page_break_before_feedback=True,
                include_edited_text_section=False,
            )

    def test_write_plain_copy_writes_paragraphs_in_order(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            out_path = Path(tmpdir) / "nested" / "copied.docx"
            paragraphs = ["First paragraph", "", " Third paragraph "]

            svc = DocxOutputService(author="Alice")
            written = svc.write_plain_copy(output_path=out_path, paragraphs=paragraphs)

            self.assertEqual(written, out_path)
            self.assertTrue(out_path.exists())

            doc = Document(str(out_path))
            self.assertEqual([p.text for p in doc.paragraphs], paragraphs)

    def test_append_paragraphs_appends_to_existing_doc(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            out_path = Path(tmpdir) / "nested" / "copied.docx"
            svc = DocxOutputService(author="Alice")
            svc.write_plain_copy(output_path=out_path, paragraphs=["first"])

            appended = svc.append_paragraphs(output_path=out_path, paragraphs=["", "second"])

            self.assertEqual(appended, out_path)
            doc = Document(str(out_path))
            self.assertEqual([p.text for p in doc.paragraphs], ["first", "", "second"])

    def test_append_corrected_paragraph_writes_no_change_message_when_identical(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            out_path = Path(tmpdir) / "nested" / "copied.docx"
            svc = DocxOutputService(author="Alice")
            svc.write_plain_copy(output_path=out_path, paragraphs=["original"])

            appended = svc.append_corrected_paragraph(
                output_path=out_path,
                original_paragraph="Same sentence.",
                corrected_paragraph="Same sentence.",
            )

            self.assertEqual(appended, out_path)
            doc = Document(str(out_path))
            self.assertEqual([p.text for p in doc.paragraphs][-2:], ["Corrected Paragraph", "No grammar corrections necessary"])


if __name__ == "__main__":
    unittest.main()
