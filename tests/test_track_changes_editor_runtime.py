from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from docx_tools.track_changes_editor import TrackChangesEditor


class TrackChangesEditorRuntimeTests(unittest.TestCase):
    def test_split_into_sentences_handles_empty_and_basic_text(self) -> None:
        self.assertEqual(TrackChangesEditor.split_into_sentences(""), [])
        text = "One sentence. Two sentence? Three sentence!"
        self.assertEqual(
            TrackChangesEditor.split_into_sentences(text),
            ["One sentence.", "Two sentence?", "Three sentence!"],
        )

    def test_revision_id_increment_and_reset(self) -> None:
        editor = TrackChangesEditor(author="A", date="2024-01-01T00:00:00Z")
        self.assertEqual(editor.next_rev_id(), 1)
        self.assertEqual(editor.next_rev_id(), 2)
        editor.reset_rev_ids()
        self.assertEqual(editor.next_rev_id(), 1)

    def test_enable_track_revisions_inserts_once(self) -> None:
        doc = Document()
        TrackChangesEditor.enable_track_revisions(doc)
        TrackChangesEditor.enable_track_revisions(doc)
        xml = doc.settings._element.xml
        self.assertEqual(xml.count("w:trackRevisions"), 1)

    def test_apply_word_diff_emits_insert_and_delete_markup(self) -> None:
        editor = TrackChangesEditor(author="A", date="2024-01-01T00:00:00Z")
        doc = Document()
        p = doc.add_paragraph()

        editor.apply_word_diff(p, "I like apples", "I really like oranges")
        xml = p._p.xml
        self.assertIn("w:ins", xml)
        self.assertIn("w:del", xml)

    def test_apply_sentence_aligned_diff_handles_all_paths(self) -> None:
        editor = TrackChangesEditor(author="A", date="2024-01-01T00:00:00Z")
        doc = Document()
        p = doc.add_paragraph()

        original = "Keep this. Remove this. Replace this."
        edited = "Keep this. Replace that. Insert this."
        editor.apply_sentence_aligned_diff(p, original, edited)

        xml = p._p.xml
        self.assertIn("w:ins", xml)
        self.assertIn("w:del", xml)

    def test_build_single_paragraph_report_writes_expected_sections(self) -> None:
        editor = TrackChangesEditor(author="A", date="2024-01-01T00:00:00Z")
        with tempfile.TemporaryDirectory() as tmpdir:
            out = Path(tmpdir) / "single.docx"
            editor.build_single_paragraph_report(
                output_path=str(out),
                original_paragraphs=["Original paragraph."],
                edited_text="Edited text.",
                corrected_text="Corrected text.",
                feedback_paragraphs=["Feedback line"],
            )
            self.assertTrue(out.exists())
            doc = Document(str(out))
            text_blob = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("ORIGINAL TEXT", text_blob)
            self.assertIn("CORRECTED TEXT", text_blob)
            self.assertIn("Language Feedback", text_blob)

    def test_build_report_with_header_and_body_writes_header_lines(self) -> None:
        editor = TrackChangesEditor(author="A", date="2024-01-01T00:00:00Z")
        with tempfile.TemporaryDirectory() as tmpdir:
            out = Path(tmpdir) / "header_body.docx"
            editor.build_report_with_header_and_body(
                output_path=str(out),
                original_paragraphs=["Original paragraph."],
                edited_text="Edited text.",
                header_lines=["Name: Dan", "Class: ENG"],
                edited_body_text="Body before.",
                corrected_body_text="Body after.",
                feedback_paragraphs=["Feedback line"],
            )
            self.assertTrue(out.exists())
            doc = Document(str(out))
            text_blob = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("CORRECTED TEXT", text_blob)
            self.assertIn("Name: Dan", text_blob)
            self.assertIn("Class: ENG", text_blob)


if __name__ == "__main__":
    unittest.main()
