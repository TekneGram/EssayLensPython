from __future__ import annotations

import tempfile
import unittest
import warnings
from pathlib import Path

from docx import Document

from docx_tools.incremental_track_changes_editor import IncrementalTrackChangesEditor

warnings.filterwarnings(
    "ignore",
    message=r"datetime\.datetime\.utcnow\(\) is deprecated.*",
    category=DeprecationWarning,
)


class IncrementalTrackChangesEditorTests(unittest.TestCase):
    def test_collect_and_remove_feedback_sections(self) -> None:
        editor = IncrementalTrackChangesEditor(author="A", date="2024-01-01T00:00:00Z")
        doc = Document()
        editor.append_feedback_section(
            doc,
            feedback_paragraphs=["Issue one.", "## Next", "Fix it."],
        )

        blocks = editor.collect_feedback_blocks(doc)
        removed = editor.remove_feedback_sections(doc)

        self.assertEqual(len(blocks), 1)
        self.assertEqual(removed, 1)
        self.assertFalse(any("ELP_FEEDBACK_BLOCK_START::" in p.text for p in doc.paragraphs))

    def test_collect_feedback_blocks_raises_on_marker_mismatch(self) -> None:
        editor = IncrementalTrackChangesEditor(author="A", date="2024-01-01T00:00:00Z")
        doc = Document()
        doc.add_paragraph("ELP_FEEDBACK_BLOCK_START::abc")
        doc.add_paragraph("Some text")
        doc.add_paragraph("ELP_FEEDBACK_BLOCK_END::xyz")

        with self.assertRaises(RuntimeError):
            editor.collect_feedback_blocks(doc)

    def test_sync_rev_id_uses_existing_revision_ids(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "report.docx"
            editor = IncrementalTrackChangesEditor(author="A", date="2024-01-01T00:00:00Z")
            doc = Document()
            p = doc.add_paragraph()
            editor._append_tracked_insertion(p, "hello")
            editor._append_tracked_deletion(p, "world")
            editor.save(doc, output)

            reloaded = IncrementalTrackChangesEditor(author="A", date="2024-01-01T00:00:00Z")
            _ = reloaded.create_or_load_document(output)
            self.assertEqual(reloaded._next_rev_id(), 3)


if __name__ == "__main__":
    unittest.main()
