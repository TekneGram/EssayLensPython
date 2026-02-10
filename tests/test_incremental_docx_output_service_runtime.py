from __future__ import annotations

import tempfile
import unittest
import warnings
from dataclasses import dataclass
from pathlib import Path

DOCX_AVAILABLE = True
try:
    from docx import Document
    from services.incremental_docx_output_service import IncrementalDocxOutputService
    from services.incremental_docx_state_store import IncrementalDocxStateStore, StateStoreError
except ModuleNotFoundError:
    DOCX_AVAILABLE = False

warnings.filterwarnings(
    "ignore",
    message=r"datetime\.datetime\.utcnow\(\) is deprecated.*",
    category=DeprecationWarning,
)


@dataclass
class _FakeResponse:
    content: str


class _FakeLlmService:
    def chat(self, **kwargs):
        return _FakeResponse(content="Strengths: clear ideas. Priorities: fix grammar and transitions.")


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx is required for incremental DOCX output service tests.")
class IncrementalDocxOutputServiceRuntimeTests(unittest.TestCase):
    def test_append_corrected_requires_edited_baseline(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "report.docx"
            service = IncrementalDocxOutputService(author="Alice")

            with self.assertRaises(StateStoreError):
                service.append_corrected_section(
                    output_path=output,
                    edited_text="Edited draft.",
                    corrected_text="Corrected draft.",
                )

    def test_feedback_summary_append_and_replace_modes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "report.docx"
            service = IncrementalDocxOutputService(author="Alice", llm_service=_FakeLlmService())

            service.start_original_section(
                output_path=output,
                original_paragraphs=["Original paragraph."],
            )
            service.append_feedback_section(
                output_path=output,
                feedback_paragraphs=["Issue: run-on sentence.", "## Next step", "Split long clauses."],
            )

            service.append_feedback_summary(
                output_path=output,
                mode="append_summary",
            )
            doc = Document(str(output))
            all_text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Feedback Summary", all_text)
            self.assertIn("ELP_FEEDBACK_BLOCK_START::", all_text)

            service.append_feedback_summary(
                output_path=output,
                mode="replace_with_summary",
            )
            doc = Document(str(output))
            all_text_after_replace = "\n".join(p.text for p in doc.paragraphs)
            self.assertNotIn("ELP_FEEDBACK_BLOCK_START::", all_text_after_replace)

    def test_hash_validation_blocks_mismatched_corrected_append(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "report.docx"
            service = IncrementalDocxOutputService(author="Alice")

            service.append_header_and_edited_section(
                output_path=output,
                header_lines=["Student Name: Alex"],
                edited_text="Original edited text.",
            )

            with self.assertRaises(StateStoreError):
                service.append_corrected_section(
                    output_path=output,
                    edited_text="Different edited text.",
                    corrected_text="Corrected text.",
                )

            state_store = IncrementalDocxStateStore(output_path=output)
            state = state_store.load()
            self.assertIn("header_edited", state["sections_written"])


if __name__ == "__main__":
    unittest.main()
