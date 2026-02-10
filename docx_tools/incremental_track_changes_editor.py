from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
import re
import uuid

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_tools.incremental_sentence_diff import align_sentences, diff_words

_FEEDBACK_START_PREFIX = "ELP_FEEDBACK_BLOCK_START::"
_FEEDBACK_END_PREFIX = "ELP_FEEDBACK_BLOCK_END::"


@dataclass
class IncrementalTrackChangesEditor:
    author: str = "EssayLens"
    date: str | None = None
    _rev_id: int = 1

    def __post_init__(self) -> None:
        if self.date is None:
            self.date = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    def create_or_load_document(self, output_path: str | Path) -> Document:
        path = Path(output_path)
        if path.exists():
            doc = Document(str(path))
        else:
            path.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
        self._sync_rev_id(doc)
        return doc

    @staticmethod
    def enable_track_revisions(doc: Document) -> None:
        settings = doc.settings._element
        if settings.find(qn("w:trackRevisions")) is None:
            settings.append(OxmlElement("w:trackRevisions"))

    def append_original_section(self, doc: Document, original_paragraphs: list[str]) -> None:
        self._add_heading(doc, "ORIGINAL TEXT")
        if not original_paragraphs:
            doc.add_paragraph("")
            return
        for paragraph in original_paragraphs:
            doc.add_paragraph(paragraph or "")

    def append_header_and_edited_section(
        self,
        doc: Document,
        header_lines: list[str],
        edited_text: str,
        include_page_break: bool = True,
    ) -> None:
        if include_page_break and doc.paragraphs:
            doc.add_page_break()
        self._add_heading(doc, "EDITED TEXT")
        for line in header_lines:
            doc.add_paragraph(line or "")
        doc.add_paragraph((edited_text or "").strip())

    def append_corrected_diff_section(self, doc: Document, edited_text: str, corrected_text: str) -> None:
        if doc.paragraphs:
            doc.add_page_break()
        self._add_heading(doc, "CORRECTED TEXT")
        paragraph = doc.add_paragraph()
        self._apply_sentence_aligned_diff(
            paragraph,
            original_text=(edited_text or "").strip(),
            edited_text=(corrected_text or "").strip(),
        )

    def append_feedback_section(
        self,
        doc: Document,
        feedback_paragraphs: list[str],
        heading: str = "Language Feedback",
    ) -> str:
        feedback_id = uuid.uuid4().hex
        marker_start = f"{_FEEDBACK_START_PREFIX}{feedback_id}"
        marker_end = f"{_FEEDBACK_END_PREFIX}{feedback_id}"

        if doc.paragraphs:
            doc.add_page_break()
        self._add_heading(doc, heading)
        doc.add_paragraph(marker_start)

        for line in feedback_paragraphs:
            text = (line or "").rstrip()
            if not text:
                doc.add_paragraph("")
                continue

            is_h2 = text.startswith("## ")
            content = text[3:].strip() if is_h2 else text.strip()
            paragraph = doc.add_paragraph()
            if is_h2:
                try:
                    paragraph.style = "Heading 2"
                except Exception:
                    pass
            paragraph.add_run(content)

        doc.add_paragraph(marker_end)
        return feedback_id

    def append_feedback_summary_section(
        self,
        doc: Document,
        feedback_summary: str,
        heading: str = "Feedback Summary",
    ) -> None:
        if doc.paragraphs:
            doc.add_page_break()
        self._add_heading(doc, heading)
        doc.add_paragraph((feedback_summary or "").strip())

    def remove_feedback_sections(self, doc: Document) -> int:
        paragraphs = list(doc.paragraphs)
        ranges: list[tuple[int, int]] = []
        open_start: tuple[int, str] | None = None

        for idx, paragraph in enumerate(paragraphs):
            text = (paragraph.text or "").strip()
            if text.startswith(_FEEDBACK_START_PREFIX):
                marker_id = text.removeprefix(_FEEDBACK_START_PREFIX)
                if open_start is not None:
                    raise RuntimeError("Feedback marker corruption: nested start marker detected.")
                open_start = (idx, marker_id)
                continue

            if text.startswith(_FEEDBACK_END_PREFIX):
                marker_id = text.removeprefix(_FEEDBACK_END_PREFIX)
                if open_start is None:
                    raise RuntimeError("Feedback marker corruption: end marker without start marker.")
                if open_start[1] != marker_id:
                    raise RuntimeError("Feedback marker corruption: marker IDs do not match.")
                ranges.append((open_start[0], idx))
                open_start = None

        if open_start is not None:
            raise RuntimeError("Feedback marker corruption: unmatched feedback start marker.")

        for start, end in reversed(ranges):
            for idx in range(end, start - 1, -1):
                paragraph = doc.paragraphs[idx]
                paragraph._p.getparent().remove(paragraph._p)

        return len(ranges)

    def collect_feedback_blocks(self, doc: Document) -> list[str]:
        blocks: list[str] = []
        current_lines: list[str] | None = None
        current_id: str | None = None

        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if text.startswith(_FEEDBACK_START_PREFIX):
                marker_id = text.removeprefix(_FEEDBACK_START_PREFIX)
                if current_lines is not None:
                    raise RuntimeError("Feedback marker corruption while collecting blocks.")
                current_lines = []
                current_id = marker_id
                continue

            if text.startswith(_FEEDBACK_END_PREFIX):
                marker_id = text.removeprefix(_FEEDBACK_END_PREFIX)
                if current_lines is None or current_id != marker_id:
                    raise RuntimeError("Feedback marker corruption while collecting blocks.")
                blocks.append("\n".join(line for line in current_lines if line))
                current_lines = None
                current_id = None
                continue

            if current_lines is not None:
                current_lines.append(paragraph.text or "")

        if current_lines is not None:
            raise RuntimeError("Feedback marker corruption: feedback block end marker missing.")

        return blocks

    def save(self, doc: Document, output_path: str | Path) -> None:
        doc.save(str(output_path))

    def _next_rev_id(self) -> int:
        rid = self._rev_id
        self._rev_id += 1
        return rid

    @staticmethod
    def _add_heading(doc: Document, title: str) -> None:
        paragraph = doc.add_paragraph()
        try:
            paragraph.style = "Heading 1"
        except Exception:
            pass
        paragraph.add_run(title)

    @staticmethod
    def _append_plain_run(paragraph, text: str) -> None:
        if not text:
            return
        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.set(qn("xml:space"), "preserve")
        text_node.text = text
        run.append(text_node)
        paragraph._p.append(run)

    def _append_tracked_insertion(self, paragraph, text: str) -> None:
        if not text:
            return
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), str(self._next_rev_id()))
        ins.set(qn("w:author"), self.author)
        ins.set(qn("w:date"), self.date)

        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.set(qn("xml:space"), "preserve")
        text_node.text = text
        run.append(text_node)
        ins.append(run)
        paragraph._p.append(ins)

    def _append_tracked_deletion(self, paragraph, text: str) -> None:
        if not text:
            return

        deletion = OxmlElement("w:del")
        deletion.set(qn("w:id"), str(self._next_rev_id()))
        deletion.set(qn("w:author"), self.author)
        deletion.set(qn("w:date"), self.date)

        run = OxmlElement("w:r")
        del_text = OxmlElement("w:delText")
        del_text.set(qn("xml:space"), "preserve")
        del_text.text = text
        run.append(del_text)
        deletion.append(run)
        paragraph._p.append(deletion)

    def _apply_word_diff(self, paragraph, original_text: str, edited_text: str) -> None:
        for op in diff_words(original_text, edited_text):
            if op.tag == "equal":
                self._append_plain_run(paragraph, op.original_text)
            elif op.tag == "delete":
                self._append_tracked_deletion(paragraph, op.original_text)
            elif op.tag == "insert":
                self._append_tracked_insertion(paragraph, op.edited_text)
            elif op.tag == "replace":
                self._append_tracked_deletion(paragraph, op.original_text)
                self._append_tracked_insertion(paragraph, op.edited_text)

    def _apply_sentence_aligned_diff(self, paragraph, original_text: str, edited_text: str) -> None:
        for tag, original_sentence, edited_sentence in align_sentences(original_text, edited_text):
            if tag == "equal" and original_sentence is not None:
                self._append_plain_run(paragraph, original_sentence + " ")
                continue
            if tag == "delete" and original_sentence is not None:
                self._append_tracked_deletion(paragraph, original_sentence + " ")
                continue
            if tag == "insert" and edited_sentence is not None:
                self._append_tracked_insertion(paragraph, edited_sentence + " ")
                continue
            if tag == "replace" and original_sentence is not None and edited_sentence is not None:
                self._apply_word_diff(paragraph, original_sentence + " ", edited_sentence + " ")

    def _sync_rev_id(self, doc: Document) -> None:
        max_id = 0
        revision_tags = {
            qn("w:ins"),
            qn("w:del"),
        }
        for element in doc.element.iter():
            if element.tag not in revision_tags:
                continue
            value = element.get(qn("w:id"))
            if not value:
                continue
            if re.fullmatch(r"\d+", value):
                max_id = max(max_id, int(value))
        self._rev_id = max_id + 1
