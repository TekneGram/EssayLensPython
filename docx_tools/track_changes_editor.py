from __future__ import annotations

import re
import difflib
from dataclasses import dataclass
from datetime import datetime
from typing import List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@dataclass
class TrackChangesEditor:
    """
    Emits Word track-changes markup (<w:ins>, <w:del>) into a NEW output .docx.
    """
    author: str = "EssayLens"
    date: Optional[str] = None
    _rev_id: int = 1

    _sentence_endings = re.compile(r"(?<=[.!?])\s+")

    def __post_init__(self) -> None:
        if self.date is None:
            self.date = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")

    @classmethod
    def split_into_sentences(cls, text: str) -> List[str]:
        txt = (text or "").strip()
        if not txt:
            return []
        return cls._sentence_endings.split(txt)

    def reset_rev_ids(self) -> None:
        self._rev_id = 1

    def next_rev_id(self) -> int:
        rid = self._rev_id
        self._rev_id += 1
        return rid

    @staticmethod
    def enable_track_revisions(doc: Document) -> None:
        settings = doc.settings._element
        if settings.find(qn("w:trackRevisions")) is None:
            settings.append(OxmlElement("w:trackRevisions"))

    @staticmethod
    def append_plain_run(paragraph, text: str) -> None:
        if not text:
            return
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        paragraph._p.append(r)

    def add_tracked_insertion(self, paragraph, text: str) -> None:
        if not text:
            return
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), str(self.next_rev_id()))
        ins.set(qn("w:author"), self.author)
        ins.set(qn("w:date"), self.date)

        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)

        ins.append(r)
        paragraph._p.append(ins)

    def add_tracked_deletion(self, paragraph, text: str) -> None:
        if not text or not text.strip():
            return

        delete = OxmlElement("w:del")
        delete.set(qn("w:id"), str(self.next_rev_id()))
        delete.set(qn("w:author"), self.author)
        delete.set(qn("w:date"), self.date)

        r = OxmlElement("w:r")
        del_text = OxmlElement("w:delText")
        del_text.set(qn("xml:space"), "preserve")
        del_text.text = text
        r.append(del_text)

        delete.append(r)
        paragraph._p.append(delete)

    def apply_word_diff(self, paragraph, original: str, edited: str) -> None:
        orig_tokens = (original or "").split()
        edit_tokens = (edited or "").split()
        matcher = difflib.SequenceMatcher(None, orig_tokens, edit_tokens)

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == "equal":
                self.append_plain_run(paragraph, " ".join(orig_tokens[i1:i2]) + " ")
            elif tag == "delete":
                self.add_tracked_deletion(paragraph, " ".join(orig_tokens[i1:i2]) + " ")
            elif tag == "insert":
                self.add_tracked_insertion(paragraph, " ".join(edit_tokens[j1:j2]) + " ")
            elif tag == "replace":
                self.add_tracked_deletion(paragraph, " ".join(orig_tokens[i1:i2]) + " ")
                self.add_tracked_insertion(paragraph, " ".join(edit_tokens[j1:j2]) + " ")

    def apply_sentence_aligned_diff(self, paragraph, original_text: str, edited_text: str) -> None:
        original_sentences = self.split_into_sentences(original_text)
        edited_sentences = self.split_into_sentences(edited_text)

        sent_matcher = difflib.SequenceMatcher(None, original_sentences, edited_sentences)

        for tag, i1, i2, j1, j2 in sent_matcher.get_opcodes():
            if tag == "equal":
                for s in original_sentences[i1:i2]:
                    self.append_plain_run(paragraph, s + " ")
            elif tag == "delete":
                for s in original_sentences[i1:i2]:
                    self.add_tracked_deletion(paragraph, s + " ")
            elif tag == "insert":
                for s in edited_sentences[j1:j2]:
                    self.add_tracked_insertion(paragraph, s + " ")
            elif tag == "replace":
                pairs = min(i2 - i1, j2 - j1)
                for k in range(pairs):
                    self.apply_word_diff(paragraph, original_sentences[i1 + k], edited_sentences[j1 + k])
                for s in original_sentences[i1 + pairs:i2]:
                    self.add_tracked_deletion(paragraph, s + " ")
                for s in edited_sentences[j1 + pairs:j2]:
                    self.add_tracked_insertion(paragraph, s + " ")

    def build_single_paragraph_report(
        self,
        output_path: str,
        original_paragraphs: List[str],
        edited_text: str,
        corrected_text: str,
        feedback_heading: str = "Language Feedback",
        feedback_paragraphs: Optional[List[str]] = None,
        feedback_as_tracked_insertion: bool = False,
        add_page_break_before_feedback: bool = True,
        include_edited_text_section: bool = True,
    ) -> None:
        self.reset_rev_ids()

        out_doc = Document()
        self.enable_track_revisions(out_doc)

        def add_h1(title: str) -> None:
            p = out_doc.add_paragraph()
            try:
                p.style = "Heading 1"
            except Exception:
                pass
            p.add_run(title)

        add_h1("ORIGINAL TEXT")
        if original_paragraphs:
            for ptxt in original_paragraphs:
                out_doc.add_paragraph(ptxt or "")
        else:
            out_doc.add_paragraph("")

        out_doc.add_page_break()

        if include_edited_text_section:
            add_h1("EDITED TEXT")
            out_doc.add_paragraph((edited_text or "").strip())
            out_doc.add_page_break()

        add_h1("CORRECTED TEXT")
        diff_p = out_doc.add_paragraph()
        self.apply_sentence_aligned_diff(
            diff_p,
            (edited_text or "").strip(),
            (corrected_text or "").strip(),
        )

        if feedback_paragraphs:
            if add_page_break_before_feedback:
                out_doc.add_page_break()

            add_h1(feedback_heading)

            for line in feedback_paragraphs:
                line = (line or "").rstrip()
                if not line.strip():
                    out_doc.add_paragraph("")
                    continue

                is_h2 = line.startswith("## ")
                text = line[3:].strip() if is_h2 else line.strip()

                p = out_doc.add_paragraph()
                if is_h2:
                    try:
                        p.style = "Heading 2"
                    except Exception:
                        pass

                if feedback_as_tracked_insertion:
                    self.add_tracked_insertion(p, text)
                else:
                    p.add_run(text)

        out_doc.save(output_path)

    def build_report_with_header_and_body(
        self,
        output_path: str,
        original_paragraphs: List[str],
        edited_text: str,
        header_lines: List[str],
        edited_body_text: str,
        corrected_body_text: str,
        feedback_heading: str = "Language Feedback",
        feedback_paragraphs: Optional[List[str]] = None,
        feedback_as_tracked_insertion: bool = False,
        add_page_break_before_feedback: bool = True,
        include_edited_text_section: bool = True,
    ) -> None:
        self.reset_rev_ids()

        out_doc = Document()
        self.enable_track_revisions(out_doc)

        def add_h1(title: str) -> None:
            p = out_doc.add_paragraph()
            try:
                p.style = "Heading 1"
            except Exception:
                pass
            p.add_run(title)

        # ORIGINAL TEXT (raw)
        add_h1("ORIGINAL TEXT")
        if original_paragraphs:
            for ptxt in original_paragraphs:
                out_doc.add_paragraph(ptxt or "")
        else:
            out_doc.add_paragraph("")

        out_doc.add_page_break()

        # EDITED TEXT (optional)
        if include_edited_text_section:
            add_h1("EDITED TEXT")
            out_doc.add_paragraph((edited_text or "").strip())
            out_doc.add_page_break()

        # CORRECTED TEXT (header lines + tracked diff on body)
        add_h1("CORRECTED TEXT")
        for line in header_lines:
            out_doc.add_paragraph(line or "")

        diff_p = out_doc.add_paragraph()
        self.apply_sentence_aligned_diff(
            diff_p,
            (edited_body_text or "").strip(),
            (corrected_body_text or "").strip(),
        )

        # Feedback section
        if feedback_paragraphs:
            if add_page_break_before_feedback:
                out_doc.add_page_break()

            add_h1(feedback_heading)

            for line in feedback_paragraphs:
                line = (line or "").rstrip()
                if not line.strip():
                    out_doc.add_paragraph("")
                    continue

                is_h2 = line.startswith("## ")
                text = line[3:].strip() if is_h2 else line.strip()

                p = out_doc.add_paragraph()
                if is_h2:
                    try:
                        p.style = "Heading 2"
                    except Exception:
                        pass

                if feedback_as_tracked_insertion:
                    self.add_tracked_insertion(p, text)
                else:
                    p.add_run(text)

        out_doc.save(output_path)
